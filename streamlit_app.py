import os
import io
import datetime
from google import genai
from google.genai import types
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Konfigurasi Halaman Streamlit
st.set_page_config(
    page_title="Sistem Notulensi Rapat Otomatis - PA Paniai",
    page_icon="⚖️",
    layout="wide",
)

# Judul Aplikasi
st.title("⚖️ Sistem Notulensi Rapat Otomatis")
st.subheader("Pengadilan Agama Paniai - Mahkamah Agung Republik Indonesia")
st.markdown("---")

# Deskripsi Aplikasi
st.markdown(
    """
    Selamat datang di **Sistem Notulensi Rapat Otomatis**. Aplikasi ini dirancang khusus untuk mempermudah tugas
    Notulen Rapat Dinas di **Pengadilan Agama Paniai** dalam menyusun notulen rapat yang sangat akurat, formal,
    dan rapi sesuai dengan format Tata Naskah Dinas Mahkamah Agung RI.
    
    Aplikasi ini ditenagai oleh **Gemini API** menggunakan library `google-genai` generasi terbaru.
    """
)

# Inisialisasi API Key dari Sidebar
st.sidebar.header("🔑 Konfigurasi API")
api_key = st.sidebar.text_input(
    "Masukkan GEMINI_API_KEY Anda",
    value=os.getenv("GEMINI_API_KEY", ""),
    type="password",
    help="Kunci API Gemini diperlukan untuk memproses audio menggunakan AI.",
)

# Status Koneksi API
if api_key:
    st.sidebar.success("Kunci API Terdeteksi!")
else:
    st.sidebar.warning("⚠️ Kunci API Kosong. Silakan isi terlebih dahulu untuk memproses audio.")

# Deskripsi Format Tata Naskah
st.sidebar.info(
    """
    **Format Naskah Dinas:**
    - Kop Surat MA RI & PA Paniai
    - Kode Dokumen Standard (FM/AM/04/02)
    - Metadata Agenda & Pembahasan Berurutan
    - Kesimpulan Rapat Detail & Taktis
    - Tanda Tangan Mengetahui Pimpinan & Notulen
    """
)

# Panduan Instalasi di Sidebar
st.sidebar.markdown("---")
st.sidebar.markdown("### 📋 Panduan Instalasi Lokal")
st.sidebar.code(
    """
pip install streamlit google-genai python-docx streamlit-mic-recorder
    """,
    language="bash",
)

st.sidebar.markdown(
    """
**Cara Menjalankan:**
1. Simpan kode ini sebagai `app.py`
2. Jalankan perintah berikut di terminal:
   `streamlit run app.py`
3. Masukkan API Key di input di atas atau buat environment variable `GEMINI_API_KEY`.
    """
)


# Helper: Membuat Cell Tabel Berwarna dan Berbingkai Rapi di python-docx
def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


# Helper: Membuat Dokumen Word (.docx) Berdasarkan Markdown
def generate_docx_bytes(markdown_text):
    doc = Document()
    
    # Set Margin Standar (1 inci)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parsing sederhana baris per baris
    lines = markdown_text.split("\n")
    
    pimpinan_rapat = "Pimpinan Rapat/Ketua"
    notulen_rapat = "Sekretaris/Notulen"
    nip_pimpinan = "....................."
    nip_notulen = "....................."
    hari_tanggal_jam = "....................."
    tempat = "Ruang Rapat Pengadilan Agama Paniai"
    peserta = "....................."
    
    agenda_rows = []
    kesimpulan_rows = []
    state = "none"

    # Ambil metadata penting untuk Word layout
    for line in lines:
        trimmed = line.strip()
        if trimmed.startswith("Hari/Tanggal/Jam"):
            hari_tanggal_jam = trimmed.split(":")[1].strip() if ":" in trimmed else hari_tanggal_jam
        elif trimmed.startswith("Tempat"):
            tempat = trimmed.split(":")[1].strip() if ":" in trimmed else tempat
        elif trimmed.startswith("Pimpinan Rapat"):
            pimpinan_rapat = trimmed.split(":")[1].strip() if ":" in trimmed else pimpinan_rapat
        elif trimmed.startswith("Peserta Rapat"):
            peserta = trimmed.split(":")[1].strip() if ":" in trimmed else peserta

        if "agenda rapat" in trimmed.lower():
            state = "agenda"
            continue
        elif "kesimpulan rapat" in trimmed.lower() or "kesimpulan rapat sebagai berikut" in trimmed.lower():
            state = "kesimpulan"
            continue
        elif trimmed.startswith("---") or trimmed.startswith("===") or trimmed.startswith("Mengetahui"):
            state = "none"

        if state == "agenda":
            if trimmed and not trimmed.startswith("-") and not trimmed.startswith("="):
                agenda_rows.append(trimmed)
        elif state == "kesimpulan":
            if trimmed and not trimmed.startswith("-") and not trimmed.startswith("="):
                kesimpulan_rows.append(trimmed)

    # Ekstrak Nama & NIP di bagian akhir jika tersedia
    signature_lines = [l.strip() for l in lines if l.strip()]
    name_lines = []
    nip_lines = []
    for sl in signature_lines:
        if sl.startswith("NIP."):
            nip_lines.append(sl)
        elif "Pimpinan Rapat" not in sl and "Notulen Rapat" not in sl and "Mengetahui" not in sl and "=" not in sl and "-" not in sl and not sl.startswith("|") and ":" not in sl and len(sl) > 3:
            # Kemungkinan nama pimpinan / notulen di baris paling bawah
            if sl != pimpinan_rapat and sl != notulen_rapat:
                name_lines.append(sl)

    # Tebak nama fiktif/ekstrak
    if len(name_lines) >= 1:
        parts = [p.strip() for p in name_lines[-1].split("   ") if p.strip()]
        if len(parts) >= 1:
            pimpinan_rapat = parts[0].replace("[", "").replace("]", "")
        if len(parts) >= 2:
            notulen_rapat = parts[1].replace("[", "").replace("]", "")

    if len(nip_lines) >= 1:
        parts = [p.replace("NIP.", "").strip() for p in nip_lines[0].split("   ") if p.strip()]
        if len(parts) >= 1:
            nip_pimpinan = parts[0].replace("[", "").replace("]", "")
        if len(parts) >= 2:
            nip_notulen = parts[1].replace("[", "").replace("]", "")

    # 1. KOP SURAT
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MAHKAMAH AGUNG REPUBLIK INDONESIA\n")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run("DIREKTORAT JENDERAL BADAN PERADILAN AGAMA\n")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    run = p.add_run("PENGADILAN TINGGI AGAMA JAYAPURA\n")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = True

    run = p.add_run("PENGADILAN AGAMA PANIAI\n")
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True

    run = p.add_run("Kompleks Kantor Bupati Paniai, Paniai Timur, Paniai, Telp. 085244544676\n")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True

    run = p.add_run("www.pa-paniai.go.id, pengadilan.agama.paniai@gmail.com\n")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True

    # Pembatas Garis Ganda
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("=========================================================================")
    run_line.bold = True
    run_line.font.size = Pt(10)

    # 2. JUDUL NOTULEN
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run("NOTULEN RAPAT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.bold = True

    # 3. TABEL KODE DOKUMEN
    table = doc.add_table(rows=2, cols=4)
    table.autofit = True
    headers = ["Kode Dokumen", "Tgl. Pembuatan", "Tgl. Revisi", "Tgl. Efektif"]
    values = ["FM/AM/04/02", "02/05/2018", ".....................", "02/05/2018"]

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "F2F2F2")
        set_cell_margins(hdr_cells[i])
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"

    # Value Row
    val_cells = table.rows[1].cells
    for i, val_text in enumerate(values):
        val_cells[i].text = val_text
        set_cell_margins(val_cells[i])
        val_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        val_cells[i].paragraphs[0].runs[0].font.name = "Arial"

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 4. METADATA DETAIL
    def add_meta_paragraph(label, val):
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(6)
        r_lbl = p_meta.add_run(f"{label:<25}")
        r_lbl.bold = True
        r_lbl.font.name = 'Arial'
        r_lbl.font.size = Pt(11)
        r_val = p_meta.add_run(f": {val}")
        r_val.font.name = 'Arial'
        r_val.font.size = Pt(11)

    add_meta_paragraph("Hari/Tanggal/Jam", hari_tanggal_jam)
    add_meta_paragraph("Tempat", tempat)
    add_meta_paragraph("Pimpinan Rapat", pimpinan_rapat)
    add_meta_paragraph("Peserta Rapat", peserta)

    # Separator
    p_sep = doc.add_paragraph()
    r_sep = p_sep.add_run("-" * 80)
    r_sep.font.color.rgb = RGBColor(128, 128, 128)

    # 5. AGENDA RAPAT
    p_ag_h = doc.add_paragraph()
    p_ag_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ag_h = p_ag_h.add_run("Agenda Rapat")
    r_ag_h.bold = True
    r_ag_h.font.size = Pt(12)
    r_ag_h.font.name = "Arial"

    if agenda_rows:
        for r in agenda_rows:
            p_row = doc.add_paragraph()
            r_row = p_row.add_run(r)
            r_row.font.size = Pt(11)
            r_row.font.name = "Arial"
    else:
        # Default text if failed to parse
        p_row1 = doc.add_paragraph()
        r1 = p_row1.add_run('Rapat dibuka oleh Sekretaris PA Paniai dengan bersama-sama membaca "Bismillahirrahmanirrahim"')
        r1.font.size = Pt(11)
        r1.font.name = "Arial"
        p_row2 = doc.add_paragraph()
        r2 = p_row2.add_run('Selanjutnya rapat dipimpin oleh Sekretaris Pengadilan agama Paniai, Pembahasan Rapat dimulai dengan mendengarkan penyampaian dari masing-masing sub bagian.')
        r2.font.size = Pt(11)
        r2.font.name = "Arial"

    # Separator
    p_sep2 = doc.add_paragraph()
    r_sep2 = p_sep2.add_run("-" * 80)
    r_sep2.font.color.rgb = RGBColor(128, 128, 128)

    # 6. KESIMPULAN
    p_ks_h = doc.add_paragraph()
    p_ks_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ks_h = p_ks_h.add_run("Kesimpulan / Keputusan Rapat")
    r_ks_h.bold = True
    r_ks_h.font.size = Pt(12)
    r_ks_h.font.name = "Arial"

    if kesimpulan_rows:
        for r in kesimpulan_rows:
            p_row = doc.add_paragraph()
            r_row = p_row.add_run(r)
            r_row.font.size = Pt(11)
            r_row.font.name = "Arial"
    else:
        p_row = doc.add_paragraph()
        r_row = p_row.add_run("Poin kesimpulan rapat belum terekstrak dengan lengkap.")
        r_row.font.size = Pt(11)
        r_row.font.italic = True
        r_row.font.name = "Arial"

    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(12)
    r_close = p_close.add_run('Selanjutnya pimpinan rapat menutup rapat selanjutnya rapat ditutup dengan ucapan "ALHAMDULILLAHIRABBIL\'ALAMIN"')
    r_close.font.size = Pt(11)
    r_close.font.name = "Arial"

    # Separator
    p_sep3 = doc.add_paragraph()
    r_sep3 = p_sep3.add_run("-" * 80)
    r_sep3.font.color.rgb = RGBColor(128, 128, 128)

    # 7. TANDA TANGAN (Mengetahui)
    doc.add_paragraph("Mengetahui,")
    
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    
    # Borderless signature table
    tblPr = sig_table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    c0 = sig_table.rows[0].cells[0]
    p_c0 = c0.paragraphs[0]
    r_c0 = p_c0.add_run("Pimpinan Rapat")
    r_c0.bold = True
    r_c0.font.name = "Arial"
    r_c0.font.size = Pt(11)
    
    # Berikan spasi untuk tanda tangan
    for _ in range(4):
        p_c0 = c0.add_paragraph()
    
    r_name0 = p_c0.add_run(pimpinan_rapat)
    r_name0.bold = True
    r_name0.font.name = "Arial"
    r_name0.font.size = Pt(11)
    p_nip0 = c0.add_paragraph()
    r_nip0 = p_nip0.add_run(f"NIP. {nip_pimpinan}")
    r_nip0.font.name = "Arial"
    r_nip0.font.size = Pt(10)

    c1 = sig_table.rows[0].cells[1]
    p_c1 = c1.paragraphs[0]
    r_c1 = p_c1.add_run("Notulen Rapat")
    r_c1.bold = True
    r_c1.font.name = "Arial"
    r_c1.font.size = Pt(11)
    
    for _ in range(4):
        p_c1 = c1.add_paragraph()
        
    r_name1 = p_c1.add_run(notulen_rapat)
    r_name1.bold = True
    r_name1.font.name = "Arial"
    r_name1.font.size = Pt(11)
    p_nip1 = c1.add_paragraph()
    r_nip1 = p_nip1.add_run(f"NIP. {nip_notulen}")
    r_nip1.font.name = "Arial"
    r_nip1.font.size = Pt(10)

    # Save to BytesIO
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io.read()


# Area Layout Kolom Kiri (Input) dan Kanan (Hasil)
col_input, col_output = st.columns([1, 1])

with col_input:
    st.header("📂 Sumber Audio Rapat")
    st.write("Silakan rekam suara rapat secara langsung atau unggah file rekaman suara yang sudah ada.")

    # Pilihan Input
    input_method = st.radio("Pilih Metode Input Audio:", ("Unggah File Rekaman", "Rekam Suara Langsung"))

    audio_bytes = None
    file_name = "rapat"
    mime_type = "audio/wav"

    if input_method == "Unggah File Rekaman":
        uploaded_file = st.file_uploader("Unggah File Audio (.wav, .mp3)", type=["wav", "mp3"])
        if uploaded_file is not None:
            audio_bytes = uploaded_file.read()
            file_name = uploaded_file.name
            mime_type = uploaded_file.type
            st.audio(audio_bytes, format=mime_type)
            st.success(f"Berhasil mengunggah: {file_name}")

    elif input_method == "Rekam Suara Langsung":
        st.write("Silakan gunakan tombol perekam di bawah:")
        
        # Streamlit mic recorder (memerlukan instalasi streamlit-mic-recorder)
        try:
            from streamlit_mic_recorder import mic_recorder
            recorded = mic_recorder(
                start_prompt="🔴 Mulai Rekam Mikrofon",
                stop_prompt="⏹️ Hentikan & Simpan",
                key="recorder"
            )
            if recorded:
                audio_bytes = recorded['bytes']
                st.audio(audio_bytes, format="audio/wav")
                st.success("Suara berhasil direkam!")
        except ImportError:
            st.error("Library `streamlit-mic-recorder` tidak terinstal. Silakan instal terlebih dahulu.")
            st.info("Alternatifnya, gunakan fitur 'Unggah File Rekaman' di atas.")

    st.markdown("---")
    
    # Tombol Kirim Ke Gemini
    if st.button("🚀 Proses Notulensi Rapat", type="primary", use_container_width=True):
        if not api_key:
            st.error("Mohon isi GEMINI_API_KEY terlebih dahulu di bagian sidebar kiri.")
        elif audio_bytes is None:
            st.error("Mohon unggah file audio atau lakukan perekaman suara terlebih dahulu.")
        else:
            with st.spinner("⏳ Sedang memproses audio dengan Gemini AI (gemini-1.5-flash)... Harap tunggu sebentar."):
                try:
                    # Inisialisasi Google GenAI SDK resmi
                    client = genai.Client(api_key=api_key)
                    
                    # Konfigurasi Prompt khusus Notulen Pengadilan Agama Paniai
                    prompt_text = """
Anda adalah seorang Notulen Rapat Profesional senior di Pengadilan Agama Paniai. 
Tugas Anda adalah mendengarkan rekaman suara rapat yang diunggah dan menyusun Notulensi Rapat Dinas yang sangat akurat, formal, dan rapi secara eksak mengikuti format tata naskah dinas instansi berikut.

Hasilkan output menggunakan format Markdown dengan struktur berikut:

MAHKAMAH AGUNG REPUBLIK INDONESIA
DIREKTORAT JENDERAL BADAN PERADILAN AGAMA
PENGADILAN TINGGI AGAMA JAYAPURA
PENGADILAN AGAMA PANIAI
Kompleks Kantor Bupati Paniai, Paniai Timur, Paniai, Telp. 085244544676
www.pa-paniai.go.id, pengadilan.agama.paniai@gmail.com
================================================================================

                                NOTULEN RAPAT

| Kode Dokumen | Tgl. Pembuatan | Tgl. Revisi | Tgl. Efektif |
| :--- | :--- | :--- | :--- |
| FM/AM/04/02 | 02/05/2018 | ..................... | 02/05/2018 |

Hari/Tanggal/Jam : [Ekstrak hari, tanggal, dan jam pelaksanaan dari audio/konteks. Jika tidak ada, gunakan tanggal hari ini secara otomatis]
Tempat           : [Ekstrak lokasi pelaksanaan, default: Ruang Rapat Pengadilan Agama Paniai]
Pimpinan Rapat   : [Ekstrak nama pimpinan rapat dari audio/konteks]
Peserta Rapat    : [Isi dengan jumlah peserta] Orang

--------------------------------------------------------------------------------
                                 Agenda Rapat
--------------------------------------------------------------------------------
Rapat dibuka oleh Sekretaris PA Paniai dengan bersama-sama membaca "Bismillahirrahmanirrahim"
Selanjutnya rapat dipimpin oleh Sekretaris Pengadilan agama Paniai, Pembahasan Rapat dimulai dengan mendengarkan penyampaian dari masing-masing sub bagian, yaitu:
1. [Tuliskan poin pembahasan sub bagian/peserta secara berurutan berdasarkan isi audio secara formal, runut, dan mendalam]
2. [Poin pembahasan selanjutnya...]
3. [Dst...]

Selanjutnya kesimpulan rapat sebagai berikut:
1. [Tuliskan poin kesimpulan/keputusan rapat secara formal, tegas, dan detail]
2. [Poin kesimpulan 2...]
3. [Dst...]

Selanjutnya pimpinan rapat menutup rapat selanjutnya rapat ditutup dengan ucapan "ALHAMDULILLAHIRABBIL'ALAMIN"

--------------------------------------------------------------------------------
Mengetahui,
Pimpinan Rapat                                        Notulen Rapat


[Nama Pimpinan Rapat/Default]                         [Nama Notulen Rapat/Default]
NIP. [NIP Pimpinan]                                   NIP. [NIP Notulen]

Aturan Penting:
1. Seluruh bahasa harus menggunakan Bahasa Indonesia formal, baku, dan sesuai dengan Tata Naskah Dinas Mahkamah Agung RI.
2. Ekstrak data-data (Hari/Tanggal, Pimpinan Rapat, Nama, NIP) seakurat mungkin dari audio. Jika tidak disebut secara spesifik dalam audio, berikan placeholder logis atau nama-nama fiktif khas instansi peradilan (misal Pimpinan: H. Ahmad, S.H., M.H., Notulen: Sarah, S.Kom.) dan berikan NIP yang realistis (18 digit).
3. Buat rincian Agenda Rapat dan Pembahasan Sub-bagian secara detail, mendalam, dan profesional, menangkap seluruh substansi diskusi, masalah yang dibahas, dan usulan solusi.
4. Tulis Kesimpulan Rapat secara butir per butir (bulleted) yang aplikatif dan taktis.
5. Harap pertahankan separator garis pembatas '---' atau '===' secara persis sesuai format.
"""

                    # Kirim audio sebagai bytes menggunakan SDK google-genai
                    response = client.models.generate_content(
                        model='gemini-1.5-flash',
                        contents=[
                            types.Part.from_bytes(
                                data=audio_bytes,
                                mime_type=mime_type,
                            ),
                            prompt_text
                        ]
                    )
                    
                    st.session_state['notulensi_markdown'] = response.text
                    st.success("🎉 Notulensi Rapat berhasil diproses!")
                except Exception as e:
                    st.error(f"Terjadi kesalahan saat memproses ke Gemini API: {str(e)}")

with col_output:
    st.header("📄 Hasil Notulensi Rapat Dinas")
    
    if 'notulensi_markdown' in st.session_state:
        markdown_text = st.session_state['notulensi_markdown']
        
        # Tampilkan Markdown di Layar
        st.markdown(markdown_text)
        
        st.markdown("---")
        st.subheader("📥 Ekspor Dokumen")
        
        col_dl1, col_dl2 = st.columns(2)
        
        # Ekspor TXT
        with col_dl1:
            st.download_button(
                label="💾 Unduh sebagai File .TXT",
                data=markdown_text,
                file_name="notulen_rapat.txt",
                mime="text/plain",
                use_container_width=True,
            )
            
        # Ekspor DOCX
        with col_dl2:
            try:
                docx_bytes = generate_docx_bytes(markdown_text)
                st.download_button(
                    label="📝 Unduh sebagai Dokumen .DOCX",
                    data=docx_bytes,
                    file_name="notulen_rapat.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Gagal mempersiapkan ekspor Word: {str(e)}")
                
    else:
        st.info("Silakan unggah atau rekam audio rapat Anda di kolom sebelah kiri, lalu tekan tombol 'Proses Notulensi Rapat' untuk melihat hasilnya di sini.")
        
        # Preview Contoh Format
        st.markdown("### Contoh Tampilan Format Naskah Dinas:")
        st.code(
            """
MAHKAMAH AGUNG REPUBLIK INDONESIA
DIREKTORAT JENDERAL BADAN PERADILAN AGAMA
PENGADILAN TINGGI AGAMA JAYAPURA
PENGADILAN AGAMA PANIAI
================================================================================
                                NOTULEN RAPAT
| Kode Dokumen | Tgl. Pembuatan | Tgl. Revisi | Tgl. Efektif |
| FM/AM/04/02  | 02/05/2018     | ........... | 02/05/2018   |

Hari/Tanggal/Jam : Rabu, 15 Juli 2026 / 09.00 WIT - Selesai
Tempat           : Ruang Rapat Pengadilan Agama Paniai
Pimpinan Rapat   : H. Ahmad, S.H., M.H. (Sekretaris)
Peserta Rapat    : 15 Orang
            """,
            language="markdown"
        )
